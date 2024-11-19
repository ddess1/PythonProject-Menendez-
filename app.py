from fastapi import FastAPI, HTTPException, Depends
from sqlalchemy import create_engine, Column, Integer, String, Float, DateTime
from sqlalchemy.orm import Session, sessionmaker, declarative_base
from sqlalchemy.exc import IntegrityError
import bcrypt
from starlette.staticfiles import StaticFiles

from models import Car, Base, User
from schemas import CarCreate, CarUpdate, CarResponse, UserBase, UserCreate, UserLogin
from database import engine, SessionLocal

from typing import List, Optional

from fastapi import Security, status
from fastapi.security import OAuth2PasswordBearer
import jwt

from fastapi import FastAPI, HTTPException, Depends
from sqlalchemy.orm import Session
from models import Car, User
from schemas import CarCreate, CarResponse
from database import engine, SessionLocal
from auth import is_admin, get_current_user, create_access_token

# Инициализация FastAPI
app = FastAPI()
SECRET_KEY = "aitu"

# Указываем FastAPI обслуживать статические файлы из папки static
app.mount("/static", StaticFiles(directory="static"), name="static")


# Создание таблиц в базе данных
Base.metadata.create_all(bind=engine)

# Функция для получения сессии
def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()







# API-эндпоинт для получения данных пользователя
@app.get("/api/users/{user_id}", response_model=UserBase)
async def get_user(user_id: int):
    db = SessionLocal()
    user = db.query(User).filter(User.id == user_id).first()
    db.close()
    if user:
        return user
    return {"message": "User not found"}



@app.post("/api/register/")
async def register_user(user: UserCreate):
    with SessionLocal() as db:
        # Проверяем, существует ли пользователь с таким email
        existing_user = db.query(User).filter(User.email == user.email).first()
        if existing_user:
            raise HTTPException(status_code=400, detail="Email already registered")

        # Хэширование пароля с использованием bcrypt
        hashed_password = bcrypt.hashpw(user.password.encode('utf-8'), bcrypt.gensalt())
        new_user = User(
            name=user.name,
            email=user.email,
            password=hashed_password.decode('utf-8'),
            role=user.role if user.role else "client"  # Если роль не указана, по умолчанию "client"
        )

        db.add(new_user)
        db.commit()
        db.refresh(new_user)  # Обновить данные объекта после вставки
        return {"message": "User registered successfully", "role": new_user.role}


@app.post("/api/login/")
async def login_user(user: UserLogin):
    db = SessionLocal()
    db_user = db.query(User).filter(User.email == user.email).first()
    db.close()

    if not db_user or not bcrypt.checkpw(user.password.encode('utf-8'), db_user.password.encode('utf-8')):
        raise HTTPException(status_code=401, detail="Invalid email or password")

    # Генерация токена
    access_token = create_access_token(data={"user_id": db_user.id})
    return {"access_token": access_token, "token_type": "bearer"}











#------------------------------------------------------------------------------------------------
#------------------------------------------------------------------------------------------------
# Шаг 3: Управление пользователями
# ●	Цель: Дать пользователям возможность управлять профилем.
# ●	Действия:
# 1.	Добавь модель профиля пользователя:
# ■	id, name, email, is_admin, created_at.
# 2.	Эндпоинты:
# ■	GET /profile: Получить информацию о профиле (только для авторизованного пользователя).
# ■	PUT /profile: Редактирование информации о профиле.
# ■	DELETE /profile: Удаление профиля


# Эндпоинт для получения профиля пользователя
@app.get("/profile", response_model=UserBase)
async def get_profile(current_user: User = Depends(get_current_user)):
    return current_user
# Эндпоинт для редактирования профиля
@app.put("/profile", response_model=UserBase)
async def update_profile(user: UserCreate, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    # Получаем актуальную запись пользователя из базы данных
    db_user = db.query(User).filter(User.id == current_user.id).first()

    if not db_user:
        raise HTTPException(status_code=404, detail="User not found")

    # Обновляем только переданные поля
    db_user.name = user.name if user.name else db_user.name
    db_user.email = user.email if user.email else db_user.email
    if user.password:
        # Хэшируем новый пароль
        db_user.password = bcrypt.hashpw(user.password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

    db.commit()
    db.refresh(db_user)
    return db_user


# Эндпоинт для удаления профиля пользователя
@app.delete("/profile")
async def delete_profile(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    # Заново загружаем пользователя, чтобы он был привязан к текущей сессии
    db_user = db.query(User).filter(User.id == current_user.id).first()

    if not db_user:
        raise HTTPException(status_code=404, detail="User not found")

    db.delete(db_user)
    db.commit()
    return {"message": "Profile deleted successfully"}
#------------------------------------------------------------------------------------------------
#------------------------------------------------------------------------------------------------



#************************************************************************************************
#************************************************************************************************
#************************************************************************************************

from fastapi import UploadFile, File, HTTPException
import os
import uuid
from pathlib import Path
from models import Car
from sqlalchemy.orm import Session
from fastapi import Depends

UPLOADS_DIR = Path("static/uploads")
UPLOADS_DIR.mkdir(parents=True, exist_ok=True)  # Создаем папку для загрузки, если ее нет

@app.post("/upload/")
async def upload_car_image(car_id: int, file: UploadFile = File(...), db: Session = Depends(get_db), current_user: User = Depends(get_current_user)):
    if current_user.role != "admin":
        raise HTTPException(status_code=403, detail="You are not authorized to uploads files.")

    # Генерация уникального имени для файла
    file_extension = file.filename.split('.')[-1]
    unique_filename = f"{uuid.uuid4()}.{file_extension}"
    file_path = UPLOADS_DIR / unique_filename

    # Сохранение файла
    with open(file_path, "wb") as buffer:
        buffer.write(await file.read())

    # Обновление записи в базе данных
    car = db.query(Car).filter(Car.id == car_id).first()
    if not car:
        raise HTTPException(status_code=404, detail="Car not found")

    car.image_url = f"uploads/{unique_filename}"
    db.commit()
    db.refresh(car)

    return {"message": "File uploaded successfully", "image_url": car.image_url}
#************************************************************************************************
#************************************************************************************************
#************************************************************************************************
from fastapi.responses import FileResponse
from docx import Document
from io import BytesIO

import os
from fastapi.responses import FileResponse
from docx import Document


from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

@app.get("/download/{car_id}")
async def download_car_info(car_id: int, db: Session = Depends(get_db)):
    # Получаем информацию о машине
    car = db.query(Car).filter(Car.id == car_id).first()
    if not car:
        raise HTTPException(status_code=404, detail="Car not found")

    # Генерация документа
    document = Document()

    # Заголовок документа
    title = document.add_heading(level=1)
    run = title.add_run(f"Exclusive Car Information\n")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)  # Темно-синий
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Введение
    intro = document.add_paragraph()
    intro.add_run("Welcome to Menendez Luxury Car Dealership!\n").bold = True
    intro.add_run(
        "We take pride in offering some of the most exclusive and high-end vehicles available in the market. Below is the detailed information about the selected car from our inventory."
    )
    intro.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.add_paragraph("\n")  # Пустая строка

    # Информация о машине
    document.add_heading(f"Car Details", level=2)
    details = document.add_paragraph()
    details.add_run(f"Model: ").bold = True
    details.add_run(f"{car.name}\n")
    details.add_run(f"Brand: ").bold = True
    details.add_run(f"{car.brand}\n")
    details.add_run(f"Price: ").bold = True
    details.add_run(f"${car.price:,.2f}\n")
    details.add_run(f"Description: ").bold = True
    details.add_run(f"{car.description}\n")

    # Дополнительный текст о магазине
    document.add_paragraph("\n")  # Пустая строка
    document.add_heading("Why Choose Menendez?", level=2)
    about = document.add_paragraph()
    about.add_run(
        "Menendez Luxury Car Dealership is committed to delivering the highest quality vehicles and unparalleled customer service. "
        "We believe that purchasing a car should be an extraordinary experience, which is why we offer tailored services to meet your needs.\n\n"
    )
    about.add_run(
        "Every car in our collection is meticulously inspected and selected to ensure it meets our standards of luxury, performance, and reliability. "
        "Thank you for choosing Menendez. Your satisfaction is our top priority."
    )

    # Заключение
    document.add_paragraph("\n")  # Пустая строка
    closing = document.add_paragraph()
    closing.add_run("Visit our showroom or contact us for more information.\n").bold = True
    closing.add_run("We look forward to serving you!")
    closing.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Сохранение файла на диск
    file_name = f"{car.name.replace(' ', '_')}_info.docx"
    file_path = os.path.join("static", "uploads", file_name)
    document.save(file_path)

    # Возврат файла
    return FileResponse(
        path=file_path,
        media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        filename=file_name,
    )



#************************************************************************************************
#************************************************************************************************
#************************************************************************************************













@app.get("/cars/{car_id}", response_model=CarResponse)
async def get_car(car_id: int, db: Session = Depends(get_db)):
    car = db.query(Car).filter(Car.id == car_id).first()
    if not car:
        raise HTTPException(status_code=404, detail="Car not found")
    return car


# Эндпоинт для создания автомобиля (только для администратора)
from fastapi import Query


@app.get("/cars/", response_model=List[CarResponse])
def get_cars(
        db: Session = Depends(get_db),
        brand: Optional[str] = None,  # Параметр фильтрации по бренду
        price_min: Optional[float] = None,  # Минимальная цена
        price_max: Optional[float] = None,  # Максимальная цена
        sort_by: Optional[str] = "name",  # Параметр сортировки (по умолчанию по имени)
):
    query = db.query(Car)

    # Фильтрация по бренду
    if brand:
        query = query.filter(Car.brand == brand)

    # Фильтрация по цене
    if price_min:
        query = query.filter(Car.price >= price_min)
    if price_max:
        query = query.filter(Car.price <= price_max)

    # Сортировка
    if sort_by == "price":
        query = query.order_by(Car.price)
    else:
        query = query.order_by(Car.name)  # По умолчанию сортировка по имени

    cars = query.all()

    if not cars:
        raise HTTPException(status_code=404, detail="No cars found")

    return cars


# Эндпоинт для получения списка всех машин
@app.get("/cars/", response_model=List[CarResponse])
def get_cars(db: Session = Depends(get_db)):
    cars = db.query(Car).all()
    if not cars:
        raise HTTPException(status_code=404, detail="No cars found")
    return cars


# Эндпоинт для обновления данных о машине (только для администратора)
@app.put("/cars/{car_id}", response_model=CarResponse)
def update_car(car_id: int, car: CarUpdate, db: Session = Depends(get_db), user: dict = Depends(is_admin)):
    db_car = db.query(Car).filter(Car.id == car_id).first()
    if not db_car:
        raise HTTPException(status_code=404, detail="Car not found")

    # Обновляем только переданные поля
    for key, value in car.dict(exclude_unset=True).items():
        setattr(db_car, key, value)

    db.commit()
    db.refresh(db_car)
    return db_car



# Эндпоинт для удаления машины (только для администратора)
@app.delete("/cars/{car_id}")
def delete_car(car_id: int, db: Session = Depends(get_db), user: dict = Depends(is_admin)):
    db_car = db.query(Car).filter(Car.id == car_id).first()
    if not db_car:
        raise HTTPException(status_code=404, detail="Car not found")
    db.delete(db_car)
    db.commit()
    return {"message": "Car deleted successfully"}


@app.post("/cars/", response_model=CarResponse)
async def create_car(car: CarCreate, db: Session = Depends(get_db), user: User = Depends(get_current_user)):
    if user.role != "admin":
        raise HTTPException(status_code=status.HTTP_403_FORBIDDEN, detail="Permission denied")

    db_car = Car(
        name=car.name,
        brand=car.brand,
        price=car.price,
        description=car.description,
        image_url=car.image_url,
        video_url=car.video_url
    )

    db.add(db_car)
    db.commit()
    db.refresh(db_car)

    return db_car